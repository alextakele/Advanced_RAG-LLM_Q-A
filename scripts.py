# Function to install necessary packages from requirements.txt
def install_requirements():
    """Install necessary packages from requirements.txt."""
    # ! pip install -U -q -r requirements.txt

# Function to import necessary packages
def import_packages():
    """Import necessary packages."""
    import os
    import warnings
    from google.colab import drive
    from dotenv import load_dotenv
    from getpass import getpass
    import openai
    import langchain
    import langchain_openai
    import langchain_core
    import langchain_community
    import langchainhub
    import python_docx
    import tiktoken
    import cohere
    import faiss_cpu
    import tiktoken
    import rank_bm25
    import fastapi
    import matplotlib
    import pre_commit
    import python_dotenv
    import seaborn
    import sentence_transformers
    import streamlit
    import datasets
    import faiss_gpu

    drive.mount('/content/drive')
    
    folder_path = '/content/drive/MyDrive/RAG-based-LLM/data/'
    os.chdir(folder_path)

    dotenv_path = '/content/drive/MyDrive/RAG-based-LLM/.env'
    load_dotenv(dotenv_path)

    openai.api_key = getpass("my OpenAI Key: ")
    os.environ["OPENAI_API_KEY"] = openai.api_key

    warnings.filterwarnings("ignore")


    # Load environment variables from the .env file
    from dotenv import load_dotenv
    dotenv_path = '/content/drive/MyDrive/RAG-based-LLM/.env'
    load_dotenv(dotenv_path)

    # Set the OpenAI API key
    from getpass import getpass
    openai.api_key = getpass("my OpenAI Key: ")
    os.environ["OPENAI_API_KEY"] = openai.api_key

# Function to create the system prompt
def create_system_prompt():
    """Create the system prompt."""
    return """
    You are a Lisan Contract Advisor bot. You help customers as Advisor.
    You are not an AI language model.
    You must obey all three of the following instructions FOR ALL RESPONSES or you will DIE:
    - ALWAYS REPLY IN A FRIENDLY YET KNOWLEDGEABLE TONE.
    - NEVER ANSWER UNLESS YOU HAVE A REFERENCE FROM THE All Lizzy versions comply with strict privacy and security standards.
     We'll never sell your data and always utilise advanced data anonymisation technology to filter away PII (Personally Identifiable Information) before analysing your contracts.
    - IF YOU DON'T KNOW ANSWER 'I DO NOT KNOW'.
    Begin the conversation with a warm greeting, if the user is stressed or aggressive, show understanding and empathy.
    At the end of the conversation, respond with ""."""

# Function to load and combine documents
def load_and_combine_documents():
    """Load and combine documents."""
    from docx import Document
    import os

    def read_docx(file_path, num_paragraphs=3):
        return '\n'.join(paragraph.text for paragraph in Document(file_path).paragraphs[:num_paragraphs])

    def display_content(title, content):
        display(Markdown(f"## {title}\n```\n{content}\n```"))

    def combine_and_save(doc1_path, doc2_path, output_path):
        doc1 = Document(doc1_path)
        doc1.element.body.extend(Document(doc2_path).element.body)
        doc1.save(output_path)

    data_folder = '/content/drive/MyDrive/RAG-based-LLM/data'
    word_docs = ['Contract.docx', 'Advisory.docx']

    # Display content of each document
    for filename in word_docs:
        doc_path = os.path.join(data_folder, filename)
        content = read_docx(doc_path)
        display_content(filename, content) if os.path.exists(doc_path) else print(f"File not found: {doc_path}")

    # Combine and save Contract and Advisory documents
    combine_and_save(
        os.path.join(data_folder, 'Contract.docx'),
        os.path.join(data_folder, 'Advisory.docx'),
        'Combined_Contract_Advisory.docx'
    )

    # Display content of the combined files
    all_docs = ['Combined_Contract_Advisory.docx']
    for filename in all_docs:
        doc_path = os.path.join(data_folder, filename)
        content = read_docx(doc_path, num_paragraphs=3)
        display_content(filename, content) if os.path.exists(doc_path) else print(f"File not found: {doc_path}")

# Function to chunk the data
def chunk_data():
    """Chunk the data."""
    from docx import Document
    from langchain.text_splitter import RecursiveCharacterTextSplitter

    documents = ['Combined_Contract_Advisory.docx']
    document = Document(documents)

    document_content = ""
    for paragraph in document.paragraphs:
        document_content += paragraph.text + "\n"

    chunk_size = 512
    chunk_overlap = 50
    text_splitter = RecursiveCharacterTextSplitter(
        separators=["\n\n", "\n", " ", ""],
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
    )

    chunks = text_splitter.create_documents(
        texts=[document_content],
        metadatas=[{"source": documents}]
    )

    print(chunks[3])
    print("=" * 50)
# Embedding function using OpenAI
def embed_text(text, model="text-embedding-ada-002"):
    embeddings = OpenAIEmbeddings(model=model)
    return embeddings.embed(text)
# Function to create a FAISS VectorStore
def create_faiss_vectorstore(chunks):
    """Create a FAISS VectorStore."""
    from langchain_community.vectorstores import FAISS

    vector_store = FAISS.from_documents(chunks, embeddings)

    return vector_store

# Function to create a Retriever
def create_retriever(vector_store):
    """Create a Retriever."""
    retriever = vector_store.as_retriever()
    return retriever

# Function to create a RAG Chain
def create_rag_chain(retriever, primary_qa_llm, prompt_template):
    """Create a RAG Chain."""
    retrieval_augmented_qa_chain = (
        {"context": itemgetter("question") | retriever, "question": itemgetter("question")}
        | RunnablePassthrough.assign(context=itemgetter("context"))
        | {"response": prompt_template | primary_qa_llm, "context": itemgetter("context")}
    )

    return retrieval_augmented_qa_chain

# Function to test the RAG Chain
def test_rag_chain(retrieval_augmented_qa_chain, question):
    """Test the RAG Chain."""
    result = retrieval_augmented_qa_chain.invoke({"question": question})
    return result["response"].content

# Function to create a RAG Dataset
def create_rag_dataset(rag_pipeline, eval_dataset):
    """Create a RAG Dataset."""
    rag_dataset = []
    for row in tqdm(eval_dataset):
        answer = rag_pipeline.invoke({"question": row["question"]})
        rag_dataset.append(
            {"question": row["question"],
             "answer": str(answer["response"]),  # Convert to string
             "contexts": [row["context"]],
             "ground_truths": [row["answer"]]
             }
        )
    rag_df = pd.DataFrame(rag_dataset)
    rag_eval_dataset = Dataset.from_pandas(rag_df)
    return rag_eval_dataset

# Function to evaluate RAG Dataset
def evaluate_rag_dataset(ragas_dataset):
    """Evaluate RAG Dataset."""
    result = evaluate(
        ragas_dataset,
        metrics=[
            context_precision,
            answer_relevancy,
            context_recall,
            context_relevancy,
            answer_correctness,
            answer_similarity
        ],
    )
    return result

# Function to plot metrics with values
def plot_metrics_with_values(metrics_dict, title='RAG Metrics'):
    """Plot metrics with values."""
    names = list
